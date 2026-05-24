"""RAG example over Word and Excel files, instrumented with ContextMesh."""
from __future__ import annotations

import argparse
import json
import os
from dataclasses import asdict, dataclass
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

from contextmesh.runtime.context_audit import audit_context_candidates
from contextmesh.runtime.context_candidates import CandidateInput, record_candidate
from contextmesh.runtime.inspector import inspect_task
from contextmesh.runtime.langfuse_export import build_langfuse_export
from contextmesh.runtime.ledger import estimate_tokens, record_step
from contextmesh.runtime.otel_export import build_otel_export
from contextmesh.runtime.team_export import build_team_export
from contextmesh.storage.db import create_db_and_tables

ROOT = Path(__file__).resolve().parent


@dataclass
class Chunk:
    ref: str
    source_type: str
    text: str
    relevance_score: float
    status: str
    reason: str


QUERY = (
    "Approve enterprise vendor renewal only if EU data residency, P1 support, "
    "security controls, and auth incident trend are acceptable."
)


def _add_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def _add_workbook(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "SLA Metrics"
    ws.append(["Metric", "Current", "Target", "Status"])
    ws.append(["P1 acknowledgement minutes", 11, 15, "green"])
    ws.append(["P1 mitigation decision minutes", 48, 60, "green"])
    ws.append(["EU data residency evidence", "available", "required", "green"])
    ws.append(["Auth incidents last 30 days", 2, 5, "green"])
    ws.append(["Legacy reset failures last 30 days", 17, 5, "red"])

    usage = wb.create_sheet("Usage Risk")
    usage.append(["Signal", "Value", "Risk"])
    usage.append(["Privileged operators", 214, "high impact"])
    usage.append(["Credential rotation dependency", "yes", "high"])
    usage.append(["EU region traffic", "82%", "material"])
    usage.append(["Marketing campaign seats", 12, "irrelevant"])
    wb.save(path)


def create_office_sources(out_dir: Path) -> dict[str, Path]:
    source_dir = out_dir / "source_files"
    source_dir.mkdir(parents=True, exist_ok=True)
    files = {
        "msa": source_dir / "Acme_MSA_Renewal.docx",
        "security": source_dir / "Security_Due_Diligence.docx",
        "legacy": source_dir / "Legacy_2023_Renewal_Notes.docx",
        "metrics": source_dir / "SLA_Usage_Risk.xlsx",
    }
    _add_docx(files["msa"], "Acme MSA Renewal Brief", [
        "Acme requires EU data residency for all incident evidence and customer records.",
        "Renewal approval requires P1 support coverage and customer-safe incident notes.",
        "The support owner is Mara Singh and procurement requires a renewal recommendation.",
    ])
    _add_docx(files["security"], "Security Due Diligence Notes", [
        "Security requires encryption at rest, SSO enforcement, and redaction of raw tokens.",
        "Vendor responses must not include secrets, raw reset links, or operator identifiers.",
        "Auth incident trend should be reviewed before renewal approval.",
    ])
    _add_docx(files["legacy"], "Legacy 2023 Renewal Notes", [
        "Legacy guidance allowed US-only evidence storage for Acme.",
        "P1 support was optional in 2023 and customer updates were sent every 4 hours.",
        "This document is superseded by the current MSA renewal requirements.",
    ])
    _add_workbook(files["metrics"])
    return files


def _docx_chunks(path: Path) -> list[tuple[str, str]]:
    doc = Document(path)
    chunks: list[tuple[str, str]] = []
    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            chunks.append((f"docx:{path.name}:p{idx}", text))
    return chunks


def _xlsx_chunks(path: Path) -> list[tuple[str, str]]:
    wb = load_workbook(path, data_only=True)
    chunks: list[tuple[str, str]] = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h) for h in rows[0]]
        for row_idx, row in enumerate(rows[1:], start=2):
            values = ["" if v is None else str(v) for v in row]
            text = "; ".join(f"{h}: {v}" for h, v in zip(headers, values, strict=False))
            chunks.append((f"xlsx:{path.name}:{ws.title}:r{row_idx}", text))
    return chunks


def _score(text: str) -> float:
    terms = [
        "eu",
        "data residency",
        "p1",
        "security",
        "auth",
        "incident",
        "credential rotation",
        "privileged",
        "renewal",
        "evidence",
        "required",
        "redaction",
        "encryption",
        "sso",
        "green",
        "red",
        "high impact",
        "high",
        "material",
        "legacy reset failures",
    ]
    lower = text.lower()
    score = sum(1 for term in terms if term in lower) / len(terms)
    if "superseded" in lower or "legacy" in lower or "2023" in lower:
        score += 0.2
    return min(1.0, round(score, 2))


def retrieve_chunks(files: dict[str, Path]) -> list[Chunk]:
    raw: list[tuple[str, str]] = []
    for key in ["msa", "security", "legacy"]:
        raw.extend(_docx_chunks(files[key]))
    raw.extend(_xlsx_chunks(files["metrics"]))

    chunks: list[Chunk] = []
    for ref, text in raw:
        score = _score(text)
        status = "selected" if score >= 0.3 else "available"
        reason = "matches renewal risk question"
        source_type = ref.split(":", 1)[0]
        lower = text.lower()
        if "Legacy_2023_Renewal_Notes" in ref:
            status = "rejected"
            score = max(score, 0.45)
            reason = "stale renewal document superseded by current requirements"
        elif "superseded" in lower or "legacy guidance" in lower:
            status = "rejected"
            score = max(score, 0.45)
            reason = "stale renewal guidance superseded by current requirements"
        elif "marketing campaign" in lower:
            status = "rejected"
            reason = "irrelevant commercial usage row"
        elif source_type == "xlsx" and any(
            marker in lower
            for marker in ["status: green", "status: red", "risk: high", "risk: material"]
        ):
            status = "selected"
            score = max(score, 0.35)
            reason = "spreadsheet risk metric required for renewal decision"
        chunks.append(Chunk(
            ref=ref,
            source_type=source_type,
            text=text,
            relevance_score=score,
            status=status,
            reason=reason,
        ))
    return chunks


def build_answer(chunks: list[Chunk]) -> dict:
    selected_text = " ".join(chunk.text for chunk in chunks if chunk.status == "selected")
    blockers = [
        "legacy reset failures are above target"
        for chunk in chunks
        if "Legacy reset failures" in chunk.text and "red" in chunk.text
    ]
    approve = not blockers and "EU data residency" in selected_text
    return {
        "query": QUERY,
        "recommendation": "conditional_approve" if not approve else "approve",
        "rationale": [
            "EU data residency evidence is available.",
            "P1 acknowledgement and mitigation decision metrics are green.",
            "Security controls require token and operator-id redaction.",
            "Legacy reset failures remain above target and require a follow-up owner.",
        ],
        "conditions": [
            "Open remediation ticket for legacy reset failures.",
            "Attach customer-safe incident-note template to renewal package.",
            "Confirm evidence remains in EU region before signature.",
        ],
    }


def run_office_rag(task_id: str, out_dir: Path) -> dict:
    create_db_and_tables()
    files = create_office_sources(out_dir)
    chunks = retrieve_chunks(files)
    answer = build_answer(chunks)
    selected = [chunk for chunk in chunks if chunk.status == "selected"]
    rejected = [chunk for chunk in chunks if chunk.status == "rejected"]

    for chunk in chunks:
        record_candidate(CandidateInput(
            task_id=task_id,
            step=1,
            ref=chunk.ref,
            status=chunk.status,
            source_type=chunk.source_type,
            reason=chunk.reason,
            relevance_score=chunk.relevance_score,
            tokens_estimated=estimate_tokens(chunk.text),
        ))

    record_step(
        task_id=task_id,
        step=1,
        agent="enterprise-office-rag-agent",
        context_refs=[chunk.ref for chunk in selected],
        context_text="\n".join(chunk.text for chunk in selected),
        decision=answer["recommendation"],
        outcome="renewal answer generated",
        outcome_class="passed",
        tokens_avoided=sum(estimate_tokens(chunk.text) for chunk in rejected),
    )

    artifacts = {
        "answer": answer,
        "chunks": [asdict(chunk) | {"text": f"{estimate_tokens(chunk.text)} tokens"} for chunk in chunks],
        "inspection": inspect_task(task_id).as_dict(),
        "audit": audit_context_candidates(task_id).as_dict(),
        "langfuse": build_langfuse_export(task_id, trace_id=f"office-rag-{task_id}").as_dict(),
        "otel": build_otel_export(task_id, service_name="enterprise-office-rag-agent").as_dict(),
        "slack": build_team_export(task_id, target="slack").as_dict(),
        "jira": build_team_export(task_id, target="jira").as_dict(),
    }
    out_dir.mkdir(parents=True, exist_ok=True)
    for name, payload in artifacts.items():
        (out_dir / f"{name}.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return artifacts


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--task-id", default="office-rag-renewal")
    parser.add_argument("--out", type=Path, default=ROOT / "out")
    args = parser.parse_args()
    if "CONTEXTMESH_STATE_DIR" not in os.environ:
        os.environ["CONTEXTMESH_STATE_DIR"] = str(args.out / ".contextmesh")
    artifacts = run_office_rag(args.task_id, args.out)
    inspection = artifacts["inspection"]
    print(
        f"{args.task_id}: {artifacts['answer']['recommendation']} "
        f"quality={inspection['context_quality_score']:.2f} "
        f"selected={len(inspection['selected_context'])} "
        f"rejected={len(inspection['rejected_context'])}"
    )
    print(f"wrote artifacts to {args.out}")


if __name__ == "__main__":
    main()
